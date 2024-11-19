from openai import OpenAI
import PyPDF2
import time
import re
from docx import Document

# Configure the OpenAI API key to authenticate requests
client = OpenAI(api_key="YOUR_KEY_HERE")

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file and organizes it into paragraphs.
    - Reads all pages of the PDF sequentially.
    - Splits text into paragraphs using double newline characters ("\n\n").
    - Preserves basic formatting for easier translation.rm 

    Args:
        pdf_path (str): The path to the PDF file to extract text from.

    Returns:
        list: A list of paragraphs extracted from the PDF.
    """
    paragraphs = []
    with open(pdf_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            paragraphs.extend(page_text.split("\n\n"))
    return paragraphs

def clean_text(text):
    """
    Cleans a given text by removing non-XML-compatible characters.
    - Removes invalid or non-printable characters that could cause issues when saving to Word.
    - Ensures the text is safe for document writing.

    Args:
        text (str): The raw text to be cleaned.

    Returns:
        str: Cleaned text without invalid characters.
    """
    return re.sub(r"[^\x09\x0A\x0D\x20-\x7E\xA0-\uFFFF]", "", text)

def translate_text_with_openai(text, source_language="English", target_language="Spanish"):
    """
    Translates a given text using OpenAI's GPT-4 model.
    - Defines a role for the AI as a translator for consistent output.
    - Sends the translation request to OpenAI's API.

    Args:
        text (str): The input text to translate.
        source_language (str): The language of the input text (default is "English").
        target_language (str): The desired language for the translation (default is "Spanish").

    Returns:
        str: The translated text.
    """
    messages = [
        {
            "role": "system",
            "content": f"You are a translator. Translate text from {source_language} to {target_language}."
        },
        {
            "role": "user",
            "content": text
        }
    ]
    response = client.chat.completions.create(
        model="gpt-4",  # Uses GPT-4 for superior translation quality
        messages=messages,
        max_tokens=3000,  # Adjust based on text size and token limit
        temperature=0.3  # Low temperature for deterministic translations
    )
    return response.choices[0].message.content.strip()

def save_paragraph_to_word(paragraph, word_path):
    """
    Appends a translated paragraph to a Word document.
    - Creates a new document if the file does not exist.
    - Ensures paragraphs are appended incrementally to avoid losing progress.

    Args:
        paragraph (str): The translated paragraph to save.
        word_path (str): The path to the Word file to save the paragraph in.
    """
    try:
        # Load existing document or create a new one if the file does not exist
        doc = Document(word_path) if os.path.exists(word_path) else Document()
    except:
        doc = Document()  # Create a new document if loading fails

    # Add the cleaned paragraph and save the document
    cleaned_paragraph = clean_text(paragraph)
    doc.add_paragraph(cleaned_paragraph)
    doc.save(word_path)

def process_pdf_translation(pdf_path, word_output_path):
    """
    Translates the content of a PDF file and saves the result incrementally in a Word document.
    - Extracts text from the PDF as a list of paragraphs.
    - Translates each paragraph individually using OpenAI.
    - Saves each translation incrementally to a Word file to prevent data loss.
    - Implements pauses to comply with OpenAI's rate limits (requests per minute).

    Args:
        pdf_path (str): Path to the input PDF file.
        word_output_path (str): Path to save the translated Word document.
    """
    print("Extracting text from PDF...")
    pdf_paragraphs = extract_text_from_pdf(pdf_path)

    print("Translating text with OpenAI...")
    request_count = 0  # Tracks the number of API requests made

    for i, paragraph in enumerate(pdf_paragraphs, 1):
        if paragraph.strip():  # Skip empty paragraphs
            print(f"Translating paragraph {i}/{len(pdf_paragraphs)}...")
            try:
                # Translate the paragraph and save it
                translated_paragraph = translate_text_with_openai(paragraph)
                save_paragraph_to_word(translated_paragraph, word_output_path)
                request_count += 1

                # Pause for 60 seconds after every 60 requests to avoid hitting the rate limit
                if request_count % 60 == 0:
                    print("Reached 60 requests. Pausing for 60 seconds...")
                    time.sleep(60)

            except Exception as e:
                # Handle translation errors and add a placeholder for failed paragraphs
                print(f"Error translating paragraph {i}: {e}")
                save_paragraph_to_word("[Error translating this paragraph]", word_output_path)

    print(f"Translation complete! Word file saved at: {word_output_path}")

# Main execution flow
if __name__ == "__main__":
    import os

    # Define paths for input and output files
    pdf_file_path = "a.pdf"  # Path to the PDF file to translate
    word_output_path = "output.docx"  # Path to save the translated Word file

    # Process the PDF and save translations incrementally
    process_pdf_translation(pdf_file_path, word_output_path)
